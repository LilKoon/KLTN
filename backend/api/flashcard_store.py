import io
import json
import os
import tempfile
from datetime import datetime, timedelta
from typing import List, Optional
from uuid import UUID
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

import database, models, schemas
from api.auth import get_current_user
from llm.orchestrator import call_llm, LLMTask
from rag.chunker import extract_text_from_pdf
from error_handler import raise_app_error

router = APIRouter(prefix="/flashcards", tags=["Flashcard Storage"])


# Vietnamese-capable TTF font path candidates for ReportLab.
_PDF_FONT_CANDIDATES = [
    "C:/Windows/Fonts/arial.ttf",
    "C:/Windows/Fonts/tahoma.ttf",
    "C:/Windows/Fonts/segoeui.ttf",
    "/Library/Fonts/Arial Unicode.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
]
_PDF_FONT_NAME: Optional[str] = None


def _ensure_pdf_font() -> str:
    global _PDF_FONT_NAME
    if _PDF_FONT_NAME:
        return _PDF_FONT_NAME
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    for path in _PDF_FONT_CANDIDATES:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont("DeckFont", path))
                _PDF_FONT_NAME = "DeckFont"
                return _PDF_FONT_NAME
            except Exception:
                continue
    _PDF_FONT_NAME = "Helvetica"  # fallback (won't render Vietnamese diacritics)
    return _PDF_FONT_NAME


def _parse_json_array(text: str) -> list:
    text = text.strip()
    if text.startswith("```"):
        text = text.split("\n", 1)[1]
        text = text.rsplit("```", 1)[0]
    parsed = json.loads(text.strip())
    if isinstance(parsed, list):
        return parsed
    for v in parsed.values():
        if isinstance(v, list):
            return v
    return list(parsed.values())


def _to_vi_card(item: dict) -> dict:
    return {
        "TuVung": item.get("word", ""),
        "Nghia": item.get("meaning_vi", ""),
        "ViDuNguCanh": item.get("example", ""),
        "LoaiTu": item.get("pos", ""),
        "PhienAm": item.get("phonetic", ""),
    }


_FLASHCARD_SYSTEM = (
    "You are an English vocabulary teacher. Respond ONLY with a valid JSON array. "
    "No markdown fences, no explanation outside JSON. "
    'Each item: {"word":"...","pos":"noun/verb/adj/adv/phrase",'
    '"phonetic":"/IPA/","meaning_vi":"nghia tieng Viet ngan gon",'
    '"example":"cau vi du tu nhien"}'
)


async def _generate_cards_from_topic(topic: str, level: str = "B1", count: int = 8) -> List[dict]:
    count = max(1, min(count, 20))
    messages = [{
        "role": "user",
        "content": (
            f"Generate {count} English vocabulary flashcards for "
            f"topic '{topic}' at CEFR level {level.upper()}."
        ),
    }]
    result = await call_llm(
        task=LLMTask.FLASHCARD_GEN,
        messages=messages,
        system_prompt=_FLASHCARD_SYSTEM,
        max_tokens=2048,
        json_mode=True,
    )
    return _parse_json_array(result["text"])


async def _generate_cards_from_text(
    text: str,
    topic: str = "Document",
    count: int = 12,
    level: str = "B1",
) -> List[dict]:
    count = max(1, min(count, 30))
    messages = [{
        "role": "user",
        "content": (
            f"Extract {count} key English vocabulary words from the following text "
            f"for flashcard set '{topic}' at CEFR level {level.upper()}. "
            f"Prefer words appropriate for that level.\n\nText:\n{text[:4000]}"
        ),
    }]
    result = await call_llm(
        task=LLMTask.FLASHCARD_GEN,
        messages=messages,
        system_prompt=_FLASHCARD_SYSTEM,
        max_tokens=2048,
        json_mode=True,
    )
    return _parse_json_array(result["text"])


class GenerateFromTopicRequest(BaseModel):
    topic: str
    level: Optional[str] = "B1"
    count: Optional[int] = 8


class CustomCard(BaseModel):
    TuVung: str
    Nghia: str
    ViDuNguCanh: Optional[str] = ""
    LoaiTu: Optional[str] = ""
    PhienAm: Optional[str] = ""


class SaveCustomDeckRequest(BaseModel):
    TenBoThe: str
    cards: List[CustomCard]
    CapDo: Optional[str] = "B1"


def calculate_sm2(ef: float, interval: int, reps: int, quality: int) -> dict:
    q = [0, 2, 4, 5][quality]
    if q < 3:
        reps = 0
        interval = 1
    else:
        if reps == 0:
            interval = {5: 7, 4: 3}.get(q, 1)
        elif reps == 1:
            interval = {5: 14, 4: 7}.get(q, 3)
        else:
            interval = round(interval * ef)
        ef = max(1.3, ef + 0.1 - (5 - q) * (0.08 + (5 - q) * 0.02))
        reps += 1
    next_due = datetime.utcnow() + timedelta(days=interval)
    return {"ef": ef, "interval": interval, "reps": reps, "next_due": next_due}


@router.get("/", response_model=List[schemas.DeckResponse])
def list_decks(
    db: Session = Depends(database.get_db),
    current_user=Depends(get_current_user),
):
    decks = (
        db.query(models.BoDTheFlashcard)
        .filter(models.BoDTheFlashcard.MaNguoiDung == current_user.MaNguoiDung)
        .all()
    )
    now = datetime.utcnow()
    result = []
    for deck in decks:
        due_count = sum(
            1 for sr in deck.trang_thai_sr if sr.NextDue <= now
        )
        result.append(
            schemas.DeckResponse(
                id=deck.MaBoDe,
                topic=deck.TenBoDe,
                level=deck.CapDo,
                count=deck.SoLuongThe,
                created_at=deck.NgayTao,
                due_today=due_count,
            )
        )
    return result


def _to_list_item(deck: models.BoDTheFlashcard) -> schemas.DeckListItem:
    return schemas.DeckListItem(
        id=deck.MaBoDe,
        title=deck.TenBoDe,
        terms=deck.SoLuongThe,
        level=deck.CapDo,
        created_at=deck.NgayTao,
    )


@router.get("/decks", response_model=List[schemas.DeckListItem])
def list_my_decks(
    db: Session = Depends(database.get_db),
    current_user=Depends(get_current_user),
):
    decks = (
        db.query(models.BoDTheFlashcard)
        .filter(models.BoDTheFlashcard.MaNguoiDung == current_user.MaNguoiDung)
        .order_by(models.BoDTheFlashcard.NgayTao.desc())
        .all()
    )
    return [_to_list_item(d) for d in decks]


@router.get("/decks/{deck_id}")
def get_deck_detail(
    deck_id: UUID,
    db: Session = Depends(database.get_db),
    current_user=Depends(get_current_user),
):
    deck = db.query(models.BoDTheFlashcard).filter(
        models.BoDTheFlashcard.MaBoDe == deck_id,
    ).first()
    if not deck:
        raise HTTPException(status_code=404, detail="Deck not found")

    owner = db.query(models.NguoiDung).filter(
        models.NguoiDung.MaNguoiDung == deck.MaNguoiDung
    ).first()
    is_public = owner is not None and owner.VaiTro == "ADMIN"
    if deck.MaNguoiDung != current_user.MaNguoiDung and not is_public:
        raise HTTPException(status_code=403, detail="Bạn không có quyền xem bộ thẻ này")

    cards = [_coerce_card(c) for c in (deck.DuLieuThe or [])]
    return {
        "id": str(deck.MaBoDe),
        "title": deck.TenBoDe,
        "level": deck.CapDo,
        "terms": deck.SoLuongThe,
        "created_at": deck.NgayTao,
        "is_public": is_public,
        "cards": cards,
    }


@router.get("/decks/{deck_id}/cards")
def get_deck_cards(
    deck_id: UUID,
    db: Session = Depends(database.get_db),
    current_user=Depends(get_current_user),
):
    deck = db.query(models.BoDTheFlashcard).filter(
        models.BoDTheFlashcard.MaBoDe == deck_id,
    ).first()
    if not deck:
        raise HTTPException(status_code=404, detail="Deck not found")
    owner = db.query(models.NguoiDung).filter(
        models.NguoiDung.MaNguoiDung == deck.MaNguoiDung
    ).first()
    is_public = owner is not None and owner.VaiTro == "ADMIN"
    if deck.MaNguoiDung != current_user.MaNguoiDung and not is_public:
        raise HTTPException(status_code=403, detail="Bạn không có quyền xem bộ thẻ này")
    return [_coerce_card(c) for c in (deck.DuLieuThe or [])]


@router.get("/public-decks", response_model=List[schemas.DeckListItem])
def list_public_decks(
    db: Session = Depends(database.get_db),
    current_user=Depends(get_current_user),
):
    decks = (
        db.query(models.BoDTheFlashcard)
        .join(models.NguoiDung, models.BoDTheFlashcard.MaNguoiDung == models.NguoiDung.MaNguoiDung)
        .filter(models.NguoiDung.VaiTro == "ADMIN")
        .order_by(models.BoDTheFlashcard.NgayTao.desc())
        .all()
    )
    return [_to_list_item(d) for d in decks]


@router.post("/clone", response_model=schemas.DeckListItem)
def clone_deck(
    req: schemas.DeckCloneRequest,
    db: Session = Depends(database.get_db),
    current_user=Depends(get_current_user),
):
    source = db.query(models.BoDTheFlashcard).filter(
        models.BoDTheFlashcard.MaBoDe == req.MaBoTheGoc,
    ).first()
    if not source:
        raise HTTPException(status_code=404, detail="Source deck not found")
    owner = db.query(models.NguoiDung).filter(
        models.NguoiDung.MaNguoiDung == source.MaNguoiDung
    ).first()
    if not owner or owner.VaiTro != "ADMIN":
        raise HTTPException(status_code=403, detail="Deck is not public")

    deck = models.BoDTheFlashcard(
        MaNguoiDung=current_user.MaNguoiDung,
        TenBoDe=source.TenBoDe,
        CapDo=source.CapDo,
        SoLuongThe=source.SoLuongThe,
        DuLieuThe=source.DuLieuThe,
    )
    db.add(deck)
    db.flush()
    for i in range(source.SoLuongThe):
        db.add(models.TrangThaiSR(MaBoDe=deck.MaBoDe, IndexThe=i))
    
    # Increment download count
    source.LuotTai = (source.LuotTai or 0) + 1

    db.commit()
    db.refresh(deck)
    return _to_list_item(deck)


@router.post("/generate/text")
async def generate_flashcards_from_topic(
    req: GenerateFromTopicRequest,
    current_user=Depends(get_current_user),
):
    if not req.topic.strip():
        raise HTTPException(status_code=422, detail="topic is required")
    try:
        raw = await _generate_cards_from_topic(req.topic, req.level or "B1", req.count or 8)
    except Exception as exc:
        raise_app_error("AI_001", detail=str(exc))
    return {"flashcards": [_to_vi_card(c) for c in raw]}


@router.post("/generate/document")
async def generate_flashcards_from_document(
    file: UploadFile = File(...),
    topic: str = Form("Document"),
    count: int = Form(12),
    level: str = Form("B1"),
    current_user=Depends(get_current_user),
):
    name = (file.filename or "").lower()
    raw_bytes = await file.read()
    if not raw_bytes:
        raise HTTPException(status_code=422, detail="Empty file")

    text = ""
    if name.endswith(".pdf"):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(raw_bytes)
            tmp_path = tmp.name
        try:
            text = extract_text_from_pdf(tmp_path)
        finally:
            try:
                os.remove(tmp_path)
            except OSError:
                pass
    elif name.endswith(".txt"):
        try:
            text = raw_bytes.decode("utf-8", errors="ignore")
        except Exception:
            text = ""
    elif name.endswith(".docx"):
        raise HTTPException(status_code=415, detail="DOCX chưa được hỗ trợ. Vui lòng dùng PDF hoặc TXT.")
    else:
        raise HTTPException(status_code=415, detail="Định dạng không hỗ trợ. Chỉ chấp nhận PDF/TXT.")

    if not text.strip():
        raise HTTPException(status_code=422, detail="Không trích xuất được nội dung từ tệp.")

    try:
        raw = await _generate_cards_from_text(text, topic=topic, count=count, level=level)
    except Exception as exc:
        raise_app_error("AI_001", detail=str(exc))
    return {"flashcards": [_to_vi_card(c) for c in raw]}


@router.post("/save/custom")
def save_custom_deck(
    req: SaveCustomDeckRequest,
    db: Session = Depends(database.get_db),
    current_user=Depends(get_current_user),
):
    valid = [c for c in req.cards if c.TuVung.strip() and c.Nghia.strip()]
    if not valid:
        raise HTTPException(status_code=422, detail="Cần ít nhất 1 thẻ hợp lệ")

    cards_data = [c.model_dump() for c in valid]
    deck = models.BoDTheFlashcard(
        MaNguoiDung=current_user.MaNguoiDung,
        TenBoDe=req.TenBoThe.strip() or "Bộ Flashcard Mới",
        CapDo=(req.CapDo or "B1").upper(),
        SoLuongThe=len(cards_data),
        DuLieuThe=cards_data,
    )
    db.add(deck)
    db.flush()
    for i in range(len(cards_data)):
        db.add(models.TrangThaiSR(MaBoDe=deck.MaBoDe, IndexThe=i))
    db.commit()
    db.refresh(deck)
    return {
        "id": str(deck.MaBoDe),
        "MaBoDe": str(deck.MaBoDe),
        "title": deck.TenBoDe,
        "terms": deck.SoLuongThe,
    }


def _coerce_card(card: dict) -> dict:
    """Cards may be stored either with Vietnamese keys (custom save) or English keys (AI-generated)."""
    return {
        "TuVung": card.get("TuVung") or card.get("word") or "",
        "Nghia": card.get("Nghia") or card.get("meaning_vi") or "",
        "PhienAm": card.get("PhienAm") or card.get("phonetic") or "",
        "LoaiTu": card.get("LoaiTu") or card.get("pos") or "",
        "ViDuNguCanh": card.get("ViDuNguCanh") or card.get("example") or "",
    }


def _build_pdf(deck: models.BoDTheFlashcard) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors

    font_name = _ensure_pdf_font()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Title"], fontName=font_name, fontSize=20, spaceAfter=8)
    meta_style = ParagraphStyle("Meta", parent=styles["Normal"], fontName=font_name, fontSize=10, textColor=colors.grey, spaceAfter=14)
    h_style = ParagraphStyle("CardHead", parent=styles["Heading3"], fontName=font_name, fontSize=13, textColor=colors.HexColor("#0f766e"), spaceAfter=2)
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontName=font_name, fontSize=11, leading=15, spaceAfter=4)
    label_style = ParagraphStyle("Label", parent=styles["Normal"], fontName=font_name, fontSize=9, textColor=colors.grey)

    story = [
        Paragraph(deck.TenBoDe, title_style),
        Paragraph(f"Trình độ: {deck.CapDo} • Số thẻ: {deck.SoLuongThe} • Tạo lúc: {deck.NgayTao.strftime('%d/%m/%Y %H:%M')}", meta_style),
    ]

    cards = [_coerce_card(c) for c in (deck.DuLieuThe or [])]
    for i, c in enumerate(cards, 1):
        head = f"{i}. {c['TuVung']}"
        if c["LoaiTu"]:
            head += f"  ({c['LoaiTu']})"
        if c["PhienAm"]:
            head += f"  {c['PhienAm']}"
        story.append(Paragraph(head, h_style))
        story.append(Paragraph(f"<b>Nghĩa:</b> {c['Nghia']}", body_style))
        if c["ViDuNguCanh"]:
            story.append(Paragraph(f"<i>Ví dụ:</i> {c['ViDuNguCanh']}", body_style))
        story.append(Spacer(1, 0.25*cm))

    doc.build(story)
    return buf.getvalue()


def _build_docx(deck: models.BoDTheFlashcard) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    document = Document()
    title = document.add_heading(deck.TenBoDe, level=0)
    meta = document.add_paragraph()
    meta_run = meta.add_run(
        f"Trình độ: {deck.CapDo} • Số thẻ: {deck.SoLuongThe} • Tạo lúc: {deck.NgayTao.strftime('%d/%m/%Y %H:%M')}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    cards = [_coerce_card(c) for c in (deck.DuLieuThe or [])]
    for i, c in enumerate(cards, 1):
        h = document.add_paragraph()
        h_run = h.add_run(f"{i}. {c['TuVung']}")
        h_run.bold = True
        h_run.font.size = Pt(13)
        h_run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
        if c["LoaiTu"]:
            extra = h.add_run(f"  ({c['LoaiTu']})")
            extra.font.size = Pt(11)
            extra.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        if c["PhienAm"]:
            ph = h.add_run(f"  {c['PhienAm']}")
            ph.font.size = Pt(11)
            ph.italic = True

        document.add_paragraph().add_run(f"Nghĩa: {c['Nghia']}").font.size = Pt(11)
        if c["ViDuNguCanh"]:
            ex = document.add_paragraph().add_run(f"Ví dụ: {c['ViDuNguCanh']}")
            ex.italic = True
            ex.font.size = Pt(11)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _safe_filename(name: str, ext: str) -> str:
    keep = "".join(ch for ch in name if ch.isalnum() or ch in (" ", "_", "-"))[:80].strip() or "flashcards"
    return f"{keep}.{ext}"


@router.get("/{deck_id}/export/pdf")
def export_deck_pdf(
    deck_id: UUID,
    db: Session = Depends(database.get_db),
    current_user=Depends(get_current_user),
):
    deck = db.query(models.BoDTheFlashcard).filter(
        models.BoDTheFlashcard.MaBoDe == deck_id,
        models.BoDTheFlashcard.MaNguoiDung == current_user.MaNguoiDung,
    ).first()
    if not deck:
        raise HTTPException(status_code=404, detail="Deck not found")
    pdf_bytes = _build_pdf(deck)
    filename = _safe_filename(deck.TenBoDe, "pdf")
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{deck_id}/export/docx")
def export_deck_docx(
    deck_id: UUID,
    db: Session = Depends(database.get_db),
    current_user=Depends(get_current_user),
):
    deck = db.query(models.BoDTheFlashcard).filter(
        models.BoDTheFlashcard.MaBoDe == deck_id,
        models.BoDTheFlashcard.MaNguoiDung == current_user.MaNguoiDung,
    ).first()
    if not deck:
        raise HTTPException(status_code=404, detail="Deck not found")
    docx_bytes = _build_docx(deck)
    filename = _safe_filename(deck.TenBoDe, "docx")
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/", response_model=schemas.DeckResponse)
def create_deck(
    req: schemas.DeckCreate,
    db: Session = Depends(database.get_db),
    current_user=Depends(get_current_user),
):
    cards_data = [c.model_dump() for c in req.cards]
    deck = models.BoDTheFlashcard(
        MaNguoiDung=current_user.MaNguoiDung,
        TenBoDe=req.topic,
        CapDo=req.level,
        SoLuongThe=len(req.cards),
        DuLieuThe=cards_data,
    )
    db.add(deck)
    db.flush()
    for i in range(len(req.cards)):
        sr = models.TrangThaiSR(MaBoDe=deck.MaBoDe, IndexThe=i)
        db.add(sr)
    db.commit()
    db.refresh(deck)
    return schemas.DeckResponse(
        id=deck.MaBoDe,
        topic=deck.TenBoDe,
        level=deck.CapDo,
        count=deck.SoLuongThe,
        created_at=deck.NgayTao,
        due_today=len(req.cards),
    )


@router.delete("/{deck_id}")
def delete_deck(
    deck_id: UUID,
    db: Session = Depends(database.get_db),
    current_user=Depends(get_current_user),
):
    deck = db.query(models.BoDTheFlashcard).filter(
        models.BoDTheFlashcard.MaBoDe == deck_id,
        models.BoDTheFlashcard.MaNguoiDung == current_user.MaNguoiDung,
    ).first()
    if not deck:
        raise HTTPException(status_code=404, detail="Deck not found")
    db.delete(deck)
    db.commit()
    return {"detail": "Deleted"}


@router.get("/{deck_id}/review")
def get_due_cards(
    deck_id: UUID,
    db: Session = Depends(database.get_db),
    current_user=Depends(get_current_user),
):
    deck = db.query(models.BoDTheFlashcard).filter(
        models.BoDTheFlashcard.MaBoDe == deck_id,
        models.BoDTheFlashcard.MaNguoiDung == current_user.MaNguoiDung,
    ).first()
    if not deck:
        raise HTTPException(status_code=404, detail="Deck not found")
    now = datetime.utcnow()
    cards_data = deck.DuLieuThe
    result = []
    for sr in deck.trang_thai_sr:
        if sr.NextDue <= now:
            result.append({
                "index": sr.IndexThe,
                "card": cards_data[sr.IndexThe],
                "sr_state": {
                    "index": sr.IndexThe,
                    "ef": sr.EasinessFactor,
                    "interval": sr.Interval,
                    "reps": sr.Repetitions,
                    "next_due": sr.NextDue.isoformat(),
                },
            })
    return result


@router.post("/{deck_id}/review")
def submit_review(
    deck_id: UUID,
    req: schemas.ReviewSubmitRequest,
    db: Session = Depends(database.get_db),
    current_user=Depends(get_current_user),
):
    deck = db.query(models.BoDTheFlashcard).filter(
        models.BoDTheFlashcard.MaBoDe == deck_id,
        models.BoDTheFlashcard.MaNguoiDung == current_user.MaNguoiDung,
    ).first()
    if not deck:
        raise HTTPException(status_code=404, detail="Deck not found")

    sr_map = {sr.IndexThe: sr for sr in deck.trang_thai_sr}
    next_dates = {}
    for item in req.results:
        if item.quality < 0 or item.quality > 3:
            raise HTTPException(status_code=422, detail="quality must be 0-3")
        sr = sr_map.get(item.card_index)
        if not sr:
            continue
        result = calculate_sm2(sr.EasinessFactor, sr.Interval, sr.Repetitions, item.quality)
        sr.EasinessFactor = result["ef"]
        sr.Interval = result["interval"]
        sr.Repetitions = result["reps"]
        sr.NextDue = result["next_due"]
        next_dates[item.card_index] = result["next_due"].isoformat()
    db.commit()
    return {"updated": len(req.results), "next_review_dates": next_dates}
